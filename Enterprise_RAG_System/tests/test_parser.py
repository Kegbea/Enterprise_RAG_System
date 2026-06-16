import pytest
from pathlib import Path
from io import BytesIO
from docx import Document as DocxDocument
from app.etl.parser import ParsedPage, BaseParser, TxtParser, MarkdownParser, get_parser

FIXTURES = Path(__file__).parent / "fixtures"


class TestParsedPage:
    def test_create_parsed_page(self):
        page = ParsedPage(page_number=1, text="hello", tables=[], headings=[])
        assert page.page_number == 1
        assert page.text == "hello"
        assert page.tables == []
        assert page.headings == []


class TestTxtParser:
    def test_supported_extensions(self):
        parser = TxtParser()
        assert "txt" in parser.supported_extensions()

    def test_parse_txt_file(self):
        parser = TxtParser()
        file_bytes = (FIXTURES / "sample.txt").read_bytes()
        pages = parser.parse(file_bytes, "sample.txt")

        assert len(pages) >= 1
        assert all(isinstance(p, ParsedPage) for p in pages)
        assert pages[0].page_number is None
        assert "企业知识管理系统" in pages[0].text
        assert pages[0].tables == []

    def test_parse_empty_txt(self):
        parser = TxtParser()
        pages = parser.parse(b"", "empty.txt")
        assert len(pages) == 1
        assert pages[0].text == ""


class TestMarkdownParser:
    def test_supported_extensions(self):
        parser = MarkdownParser()
        assert "md" in parser.supported_extensions()

    def test_parse_md_headings(self):
        parser = MarkdownParser()
        file_bytes = (FIXTURES / "sample.md").read_bytes()
        pages = parser.parse(file_bytes, "sample.md")

        assert len(pages) >= 1
        assert pages[0].page_number is None
        assert len(pages[0].headings) > 0
        assert "企业知识管理系统" in " ".join(pages[0].headings)

    def test_parse_md_table(self):
        parser = MarkdownParser()
        file_bytes = (FIXTURES / "sample.md").read_bytes()
        pages = parser.parse(file_bytes, "sample.md")

        all_tables = [t for p in pages for t in p.tables]
        assert len(all_tables) >= 1
        assert "组件" in all_tables[0]
        assert "ChromaDB" in all_tables[0]

    def test_parse_md_heading_path(self):
        parser = MarkdownParser()
        content = b"# Ch1\n\n## Sec1.1\n\nSome text here.\n\n### Sec1.1.1\n\nDeeper text.\n"
        pages = parser.parse(content, "test.md")
        text = pages[0].text
        assert "Ch1" in text
        assert "Sec1.1" in text

    def test_md_no_table(self):
        parser = MarkdownParser()
        content = b"# Title\n\nJust some paragraph text.\n\nAnother paragraph.\n"
        pages = parser.parse(content, "test.md")
        assert pages[0].tables == []


def _make_docx_bytes(paragraphs: list[str], include_table: bool = False) -> bytes:
    """辅助函数：用 python-docx 生成最小 docx 文件 bytes。"""
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    if include_table:
        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"
        headers = ["列A", "列B", "列C"]
        for j, h in enumerate(headers):
            table.rows[0].cells[j].text = h
        data = [["val1", "val2", "val3"], ["val4", "val5", "val6"]]
        for i, row_data in enumerate(data):
            for j, val in enumerate(row_data):
                table.rows[i + 1].cells[j].text = val
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxParser:
    def test_supported_extensions(self):
        from app.etl.parser import DocxParser
        parser = DocxParser()
        assert "docx" in parser.supported_extensions()

    def test_parse_docx_paragraphs(self):
        from app.etl.parser import DocxParser
        parser = DocxParser()
        file_bytes = _make_docx_bytes(["第一章 概述", "这是一段测试文本。", "这是第二段。"])
        pages = parser.parse(file_bytes, "test.docx")

        assert len(pages) == 1
        assert "第一章 概述" in pages[0].text
        assert "测试文本" in pages[0].text
        assert pages[0].page_number is None

    def test_parse_docx_with_table(self):
        from app.etl.parser import DocxParser
        parser = DocxParser()
        file_bytes = _make_docx_bytes(["以下是数据表格："], include_table=True)
        pages = parser.parse(file_bytes, "test.docx")

        assert len(pages[0].tables) >= 1
        assert "列A" in pages[0].tables[0]
        assert "val1" in pages[0].tables[0]

    def test_parse_empty_docx(self):
        from app.etl.parser import DocxParser
        parser = DocxParser()
        file_bytes = _make_docx_bytes([])
        pages = parser.parse(file_bytes, "empty.docx")
        assert len(pages) >= 1


def _make_minimal_pdf_bytes(text: str) -> bytes:
    """生成包含指定文本的最小有效 PDF（用于测试）。"""
    # 使用 PDF 1.4 最小结构，包含单页和文本
    encoded_text = text.encode("utf-8")
    # 转义 PDF 字符串中的特殊字符
    safe_text = encoded_text.replace(b"\\", b"\\\\").replace(b"(", b"\\(").replace(b")", b"\\)")

    content_stream = (
        b"BT /F1 12 Tf 100 700 Td (" + safe_text + b") Tj ET"
    )

    objects = [
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj",
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj",
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Contents 4 0 R/Resources<</Font<</F1<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>>>>>>>endobj",
        b"4 0 obj<</Length " + str(len(content_stream)).encode() + b">>stream\n" + content_stream + b"\nendstream\nendobj",
    ]

    object_offsets = []
    pdf = b"%PDF-1.4\n"
    for obj in objects:
        object_offsets.append(len(pdf))
        pdf += obj + b"\n"

    xref_offset = len(pdf)
    pdf += b"xref\n"
    pdf += b"0 " + str(len(objects) + 1).encode() + b"\n"
    pdf += b"0000000000 65535 f \n"
    for offset in object_offsets:
        pdf += f"{offset:010d} 00000 n \n".encode()

    pdf += b"trailer<</Size " + str(len(objects) + 1).encode() + b"/Root 1 0 R>>\n"
    pdf += b"startxref\n" + str(xref_offset).encode() + b"\n%%EOF"

    return pdf


class TestPDFParser:
    def test_supported_extensions(self):
        from app.etl.parser import PDFParser
        parser = PDFParser()
        assert "pdf" in parser.supported_extensions()

    def test_parse_pdf_pages(self):
        from app.etl.parser import PDFParser
        parser = PDFParser()
        file_bytes = _make_minimal_pdf_bytes("Hello World")
        pages = parser.parse(file_bytes, "test.pdf")

        assert len(pages) >= 1
        assert pages[0].page_number == 1

    def test_table_to_markdown(self):
        from app.etl.parser import PDFParser
        table = [["A", "B", "C"], ["1", "2", "3"], ["4", "5", "6"]]
        result = PDFParser._table_to_markdown(table)
        assert "| A | B | C |" in result
        assert "|---|---|---|" in result
        assert "| 1 | 2 | 3 |" in result

    def test_table_to_markdown_empty(self):
        from app.etl.parser import PDFParser
        assert PDFParser._table_to_markdown([]) == ""
        assert PDFParser._table_to_markdown([[]]) == ""

    def test_table_to_markdown_none_cells(self):
        from app.etl.parser import PDFParser
        table = [["A", None, "C"], [None, "2", None]]
        result = PDFParser._table_to_markdown(table)
        assert "| A |  | C |" in result


class TestGetParser:
    def test_get_parser_by_extension(self):
        from app.etl.parser import DocxParser, PDFParser
        assert isinstance(get_parser("file.txt"), TxtParser)
        assert isinstance(get_parser("file.md"), MarkdownParser)
        assert isinstance(get_parser("file.docx"), DocxParser)
        assert isinstance(get_parser("file.pdf"), PDFParser)

    def test_get_parser_unknown_extension(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            get_parser("file.xyz")
