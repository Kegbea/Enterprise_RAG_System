"""文档解析器 — 统一接口 + 四类格式实现 + 工厂函数。

设计原则：
- 所有解析器返回统一的 ParsedPage 结构
- 表格统一输出为 Markdown Table 格式
- 通过 get_parser() 工厂函数按扩展名分发
"""

import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from io import BytesIO

import pdfplumber
from docx import Document as DocxDocument
from docx.oxml.ns import qn


@dataclass
class ParsedPage:
    """解析后的单页结构。

    表格统一为 Markdown Table 字符串，Chunker 只需识别一种格式。
    """

    page_number: int | None  # TXT/MD 为 None
    text: str = ""
    tables: list[str] = field(default_factory=list)  # Markdown table 格式
    headings: list[str] = field(default_factory=list)  # 该页检测到的标题


class BaseParser(ABC):
    """文档解析器基类。"""

    @abstractmethod
    def parse(self, file_bytes: bytes, filename: str) -> list[ParsedPage]:
        """将文件字节流解析为 ParsedPage 列表。"""
        ...

    @abstractmethod
    def supported_extensions(self) -> list[str]:
        """返回支持的文件扩展名列表（不含点号）。"""
        ...


class TxtParser(BaseParser):
    """纯文本解析器 — 全文作为单页，无页码，无表格。"""

    def supported_extensions(self) -> list[str]:
        return ["txt"]

    def parse(self, file_bytes: bytes, filename: str) -> list[ParsedPage]:
        text = file_bytes.decode("utf-8", errors="replace")
        return [ParsedPage(page_number=None, text=text)]


class MarkdownParser(BaseParser):
    """Markdown 解析器 — 识别 # 标题层级，正则提取表格。"""

    def supported_extensions(self) -> list[str]:
        return ["md"]

    def parse(self, file_bytes: bytes, filename: str) -> list[ParsedPage]:
        text = file_bytes.decode("utf-8", errors="replace")
        headings = self._extract_headings(text)
        tables = self._extract_tables(text)
        return [ParsedPage(page_number=None, text=text, tables=tables, headings=headings)]

    @staticmethod
    def _extract_headings(text: str) -> list[str]:
        headings = []
        for line in text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("#"):
                heading = stripped.lstrip("#").strip()
                headings.append(heading)
        return headings

    @staticmethod
    def _extract_tables(text: str) -> list[str]:
        """提取 Markdown 表格（以 | 开头的连续行，含分隔行）。"""
        tables = []
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith("|") and "---" not in line:
                if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
                    table_lines = [lines[i]]
                    i += 1
                    while i < len(lines):
                        table_lines.append(lines[i])
                        if i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                            i += 1
                        else:
                            break
                    tables.append("\n".join(table_lines))
            i += 1
        return tables


class DocxParser(BaseParser):
    """DOCX 解析器 — 提取段落和表格，识别 Heading 1-6 样式。"""

    def supported_extensions(self) -> list[str]:
        return ["docx"]

    def parse(self, file_bytes: bytes, filename: str) -> list[ParsedPage]:
        buf = BytesIO(file_bytes)
        doc = DocxDocument(buf)

        text_parts = []
        tables = []
        headings = []

        for element in doc.element.body:
            if element.tag == qn("w:p"):
                para = self._find_paragraph(doc, element)
                if para is not None:
                    style = para.style.name if para.style else ""
                    para_text = para.text.strip()
                    if para_text:
                        text_parts.append(para_text)
                        if style.startswith("Heading") or style.startswith("heading"):
                            headings.append(para_text)
            elif element.tag == qn("w:tbl"):
                table = self._find_table(doc, element)
                if table is not None:
                    md_table = self._table_to_markdown(table)
                    if md_table:
                        tables.append(md_table)
                        text_parts.append(f"[TABLE:{len(tables) - 1}]")

        text = "\n\n".join(text_parts)
        return [ParsedPage(page_number=None, text=text, tables=tables, headings=headings)]

    @staticmethod
    def _find_paragraph(doc, element):
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    @staticmethod
    def _find_table(doc, element):
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    @staticmethod
    def _table_to_markdown(table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if not rows:
            return ""
        if len(rows) >= 1:
            col_count = len(table.rows[0].cells)
            separator = "|" + "|".join(["---" for _ in range(col_count)]) + "|"
            rows.insert(1, separator)
        return "\n".join(rows)


class PDFParser(BaseParser):
    """PDF 解析器 — 基于 pdfplumber 按页提取文本和表格。"""

    def supported_extensions(self) -> list[str]:
        return ["pdf"]

    def parse(self, file_bytes: bytes, filename: str) -> list[ParsedPage]:
        pages = []
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                tables_raw = page.extract_tables() or []

                tables = []
                for table_idx, table in enumerate(tables_raw):
                    md_table = self._table_to_markdown(table)
                    if md_table:
                        tables.append(md_table)
                        text += f"\n[TABLE:{table_idx}]\n"

                pages.append(ParsedPage(
                    page_number=page_num,
                    text=text,
                    tables=tables,
                    headings=[],
                ))
        return pages

    @staticmethod
    def _table_to_markdown(table: list[list[str | None]]) -> str:
        """将 pdfplumber 提取的二维表格转为 Markdown Table。"""
        if not table or not table[0]:
            return ""
        rows = []
        for row in table:
            cells = [(cell or "").strip().replace("\n", " ") for cell in row]
            rows.append("| " + " | ".join(cells) + " |")
        if not rows:
            return ""
        col_count = len(table[0])
        separator = "|" + "|".join(["---" for _ in range(col_count)]) + "|"
        rows.insert(1, separator)
        return "\n".join(rows)


# ── 解析器注册表 ──
_PARSER_REGISTRY: dict[str, BaseParser] = {}


def _register(parser: BaseParser) -> BaseParser:
    for ext in parser.supported_extensions():
        _PARSER_REGISTRY[ext] = parser
    return parser


# ── 注册内置解析器 ──
_register(TxtParser())
_register(MarkdownParser())
_register(DocxParser())
_register(PDFParser())


def get_parser(filename: str) -> BaseParser:
    """根据文件名扩展名获取对应的解析器实例。

    Raises:
        ValueError: 不支持的文件类型
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parser = _PARSER_REGISTRY.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: .{ext}")
    return parser
